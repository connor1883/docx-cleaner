import docx
import sympy as sp
import re
import os


def extract_office_math(doc_path):
    doc = docx.Document(doc_path)
    equations = []

    # Iterate through paragraphs and table cells
    for paragraph in doc.paragraphs:
        if paragraph.text:
            equations.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    equations.append(cell.text.strip())

    return equations


def office_math_to_latex(office_math_expr):
    # Replace the Office Math specific notation with sympy compatible notation
    office_math_expr = office_math_expr.replace('∑', 'Sum').replace('▒', '').replace('√', 'sqrt')

    # Handle fractions
    office_math_expr = re.sub(r'(\d+)/(\d+)', r'\\frac{\1}{\2}', office_math_expr)

    # Regular expression to find and replace subscript/superscript notation
    office_math_expr = re.sub(r'_(\{[^\}]+\}|\w)', r'_\1', office_math_expr)
    office_math_expr = re.sub(r'\^(\{[^\}]+\}|\w)', r'^\1', office_math_expr)

    # Check if the expression is empty or whitespace only
    if not office_math_expr.strip():
        return "Empty or invalid expression"

    # Define the symbols used in the expression
    i, N = sp.symbols('i N')
    X_i = sp.Function('X_i')

    # Convert the expression to a sympy expression
    try:
        expr = sp.sympify(office_math_expr, locals={'X_i': X_i(i), 'Sum': sp.Sum, 'sqrt': sp.sqrt})
    except Exception as e:
        return f"Error parsing expression: {office_math_expr}. Error: {e}"

    # Convert the sympy expression to LaTeX
    latex_expr = sp.latex(expr)

    return latex_expr


def convert_docx_to_latex(doc_path):
    if not os.path.exists(doc_path):
        print(f"File not found: {doc_path}")
        return []

    equations = extract_office_math(doc_path)
    latex_equations = []

    for eq in equations:
        latex_eq = office_math_to_latex(eq)
        latex_equations.append(latex_eq)

    return latex_equations


doc_path = r'C:\Users\Connor\Desktop\Summer 2024 Econ\OfficeMathRoute\OM - Copy.docx'  # Use raw string to avoid escape character issues
latex_equations = convert_docx_to_latex(doc_path)

for latex_eq in latex_equations:
    print(latex_eq)
